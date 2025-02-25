#!/usr/bin/env python3
"""
Convert Markdown to DOCX with proper academic formatting.
Handles bullet points, images, tables, and citations consistently.
"""

import os
import sys
import argparse
import warnings
import panflute as pf
from pathlib import Path
import subprocess
import tempfile
import shutil
import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.table import WD_TABLE_ALIGNMENT

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

warnings.filterwarnings('ignore', category=UserWarning, module='docx')

class DocumentConverter:
    def __init__(self, input_file, output_file, reference_file=None):
        self.input_file = Path(input_file)
        self.output_file = Path(output_file)
        self.reference_file = Path(reference_file) if reference_file else None
        self.root_dir = self.input_file.parent
        self.media_dir = self.root_dir / "media"
        logger.info(f"Initializing conversion: {self.input_file} → {self.output_file}")
        if self.reference_file:
            logger.info(f"Using reference file: {self.reference_file}")

    def create_reference_doc(self):
        """Create and configure reference document with proper styles."""
        logger.info("Creating reference document with custom styles...")
        
        # Create default reference doc
        logger.debug("Generating default reference.docx template")
        subprocess.run(["pandoc", "-o", str(self.reference_file), "--print-default-data-file", "reference.docx"], check=True)

        # Configure styles
        logger.info("Configuring document styles...")
        doc = Document(self.reference_file)

        # Set page dimensions and margins
        logger.debug("Setting page dimensions and margins")
        for section in doc.sections:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.top_margin = Inches(0.25)
            section.bottom_margin = Inches(0.25)
            section.left_margin = Inches(0.25)
            section.right_margin = Inches(0.25)

        # Configure styles
        logger.debug("Configuring paragraph styles")
        style_count = 0
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                if hasattr(style, 'font'):
                    style.font.name = 'Times New Roman'
                    style_count += 1
                    
                    if style.name == 'Heading 1':
                        logger.debug("Configuring Heading 1 style")
                        style.font.size = Pt(16)
                        style.font.bold = True
                        style.paragraph_format.space_before = Pt(36)
                        style.paragraph_format.space_after = Pt(18)
                        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif style.name == 'Heading 2':
                        logger.debug("Configuring Heading 2 style")
                        style.font.size = Pt(14)
                        style.font.bold = True
                        style.paragraph_format.space_before = Pt(30)
                        style.paragraph_format.space_after = Pt(12)
                    elif style.name == 'Heading 3':
                        logger.debug("Configuring Heading 3 style")
                        style.font.size = Pt(12)
                        style.font.bold = True
                        style.paragraph_format.space_before = Pt(24)
                        style.paragraph_format.space_after = Pt(10)
                    elif style.name == 'Normal':
                        logger.debug("Configuring Normal style")
                        style.font.size = Pt(12)
                        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        style.paragraph_format.space_before = Pt(12)
                        style.paragraph_format.space_after = Pt(12)
                    elif style.name == 'List Bullet':
                        logger.debug("Configuring List Bullet style")
                        style.font.size = Pt(12)
                        style.paragraph_format.left_indent = Inches(0.15)
                        style.paragraph_format.first_line_indent = Inches(-0.15)
                        style.paragraph_format.space_before = Pt(0)
                        style.paragraph_format.space_after = Pt(0)
                    elif style.name == 'Caption':
                        logger.debug("Configuring Caption style")
                        style.font.size = Pt(10)
                        style.font.italic = True
                        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        style.paragraph_format.space_before = Pt(6)
                        style.paragraph_format.space_after = Pt(12)

        logger.info(f"Configured {style_count} paragraph styles")
        doc.save(self.reference_file)
        logger.info("Reference document created successfully")

    def prepare_environment(self):
        """Prepare temporary directories and files."""
        logger.info("Preparing environment...")
        
        if not self.input_file.exists():
            logger.error(f"Input file not found: {self.input_file}")
            raise FileNotFoundError(f"Input file not found: {self.input_file}")
        
        # Create media directory if it doesn't exist
        self.media_dir.mkdir(exist_ok=True)
        logger.debug(f"Media directory ready: {self.media_dir}")
        
        # Create a temporary directory for processing
        self.temp_dir = Path(tempfile.mkdtemp())
        self.temp_file = self.temp_dir / self.input_file.name
        logger.debug(f"Created temporary directory: {self.temp_dir}")

        # Create reference doc if needed
        if not self.reference_file:
            self.reference_file = self.temp_dir / "reference.docx"
            logger.info("No reference file provided, creating default")
            self.create_reference_doc()

    def process_markdown(self):
        """Process markdown file to ensure proper formatting."""
        logger.info("Processing markdown content...")
        
        with open(self.input_file, 'r') as f:
            content = f.read()
            logger.debug(f"Read {len(content)} characters from input file")

        # Count and log occurrences of different elements
        image_count = content.count('<img src=')  
        table_count = content.count('|---')
        bullet_count = len([line for line in content.split('\n') if line.strip().startswith(('•', '-', '*'))])
        
        logger.info(f"Found in document: {image_count} images, {table_count} tables, {bullet_count} bullet points")

        # Convert bullet points to asterisks
        content = content.replace('•', '*')
        content = content.replace('  -', '  *')
        
        # Process images to ensure proper rendering
        lines = content.split('\n')
        processed_lines = []
        in_div = False
        
        for line in lines:
            if '<div' in line:
                in_div = True
                continue
            elif '</div>' in line:
                in_div = False
                continue
            elif '<img' in line and 'src=' in line:
                # Extract image path and alt text
                src = line[line.find('src="')+5:line.find('" alt=')]
                alt = line[line.find('alt="')+5:line.rfind('"')]
                # Convert to markdown image syntax
                processed_lines.append(f'![{alt}]({src})')
            else:
                processed_lines.append(line)
        
        content = '\n'.join(processed_lines)
        
        with open(self.temp_file, 'w') as f:
            f.write(content)
        logger.info("Markdown processing complete")

    def apply_final_formatting(self):
        """Apply final formatting to the output document."""
        logger.info("Applying final document formatting...")
        doc = Document(self.output_file)

        # Track formatting statistics
        stats = {
            'bullet_points': 0,
            'figures': 0,
            'tables': 0,
            'table_cells': 0
        }

        # Process paragraphs
        logger.info("Processing paragraphs and applying styles...")
        for paragraph in doc.paragraphs:
            # Handle bullet points (check for both '-' and '•' symbols)
            if any(paragraph.text.strip().startswith(sym) for sym in ['-', '•', '*']):
                stats['bullet_points'] += 1
                paragraph.paragraph_format.left_indent = Inches(0.15)
                paragraph.paragraph_format.first_line_indent = Inches(-0.15)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

            # Handle figure captions
            elif paragraph.text.strip().startswith('*Figure'):
                stats['figures'] += 1
                logger.debug(f"Formatting figure caption: {paragraph.text[:50]}...")
                paragraph.style = doc.styles['Caption']
                paragraph.paragraph_format.space_before = Pt(6)
                paragraph.paragraph_format.space_after = Pt(12)
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.italic = True
                    run.font.size = Pt(10)

            # Handle table captions
            elif paragraph.text.strip().startswith('**Table'):
                paragraph.style = doc.styles['Caption']
                paragraph.paragraph_format.space_before = Pt(12)
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    if run.text.startswith('**') and run.text.endswith('**'):
                        run.text = run.text[2:-2]  # Remove markdown bold markers
                        run.font.bold = True
                        run.font.italic = False
                    run.font.size = Pt(10)

            # Handle table notes
            elif paragraph.text.strip().startswith('*Note:'):
                paragraph.style = doc.styles['Caption']
                paragraph.paragraph_format.space_before = Pt(6)
                paragraph.paragraph_format.space_after = Pt(12)
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.italic = True
                    run.font.size = Pt(10)

        # Process images
        logger.info("Processing and resizing images...")
        for shape in doc.inline_shapes:
            if shape.width > 0:
                original_width = shape.width
                desired_width = Inches(6)
                aspect_ratio = shape.height / shape.width
                shape.width = desired_width
                shape.height = int(desired_width * aspect_ratio)
                logger.debug(f"Resized image from {original_width} to {desired_width} width")

        # Process tables
        logger.info("Formatting tables...")
        for table_idx, table in enumerate(doc.tables, 1):
            stats['tables'] += 1
            stats['table_cells'] += len(table.rows) * len(table.columns)
            logger.debug(f"Processing table {table_idx}: {len(table.rows)}x{len(table.columns)}")
            
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.allow_autofit = False
            
            # Format table
            table_width = Inches(8)
            table.width = table_width
            
            # Process cells
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    # Set cell borders
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                                        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                        '</w:tcBorders>')
                    tcPr.append(tcBorders)

                    # Format text
                    for paragraph in cell.paragraphs:
                        # Center align all cells
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(10)  # Slightly larger than minimum size
                            
                            # Make header row bold
                            if row_idx == 0:
                                run.font.bold = True

        # Log formatting statistics
        logger.info("Formatting statistics:")
        logger.info(f"- Processed {stats['bullet_points']} bullet points")
        logger.info(f"- Formatted {stats['figures']} figures")
        logger.info(f"- Processed {stats['tables']} tables with {stats['table_cells']} cells")

        doc.save(self.output_file)
        logger.info("Final formatting complete")

    def run_conversion(self):
        """Execute the conversion process."""
        try:
            logger.info("Starting document conversion process...")
            self.prepare_environment()
            self.process_markdown()

            # Build resource path
            resource_path = [
                ".",
                str(self.root_dir),
                str(self.root_dir / "sources_png"),
                str(self.root_dir / "user_illustrations")
            ]
            logger.info(f"Resource paths configured: {resource_path}")

            # Build pandoc command
            cmd = [
                "pandoc",
                str(self.temp_file),
                "-f", "markdown+raw_html+raw_tex+raw_attribute+pipe_tables+grid_tables+multiline_tables+simple_tables",
                "-t", "docx",
                "--reference-doc", str(self.reference_file),
                "--wrap=preserve",
                "--standalone",
                "--resource-path", ":".join(resource_path),
                "--dpi=300",
                "--verbose",
                "-o", str(self.output_file)
            ]
            
            logger.info("Running pandoc conversion...")
            logger.debug(f"Command: {' '.join(cmd)}")

            # Run conversion
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                check=True
            )
            
            if result.stderr:
                logger.debug(f"Pandoc output: {result.stderr}")

            # Apply final formatting
            self.apply_final_formatting()

            logger.info(f"Successfully converted {self.input_file} to {self.output_file}")
            return True

        except subprocess.CalledProcessError as e:
            logger.error(f"Pandoc conversion failed: {e.stderr}")
            return False
        except Exception as e:
            logger.error(f"Unexpected error during conversion: {str(e)}")
            return False
        finally:
            self.cleanup()

    def cleanup(self):
        """Clean up temporary files."""
        logger.info("Cleaning up temporary files...")
        shutil.rmtree(self.temp_dir, ignore_errors=True)
        logger.debug(f"Removed temporary directory: {self.temp_dir}")

def main():
    parser = argparse.ArgumentParser(description="Convert Markdown to DOCX with proper academic formatting")
    parser.add_argument("input_file", help="Input Markdown file")
    parser.add_argument("output_file", help="Output DOCX file")
    parser.add_argument("--reference", help="Reference DOCX file for styles")
    parser.add_argument("--verbose", action="store_true", help="Enable verbose logging")
    args = parser.parse_args()

    if args.verbose:
        logger.setLevel(logging.DEBUG)
    
    logger.info("Starting document conversion")
    logger.info(f"Input: {args.input_file}")
    logger.info(f"Output: {args.output_file}")

    converter = DocumentConverter(
        input_file=args.input_file,
        output_file=args.output_file,
        reference_file=args.reference
    )
    
    success = converter.run_conversion()
    if success:
        logger.info("Conversion completed successfully")
    else:
        logger.error("Conversion failed")
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
